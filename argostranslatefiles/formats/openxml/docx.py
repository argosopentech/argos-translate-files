from argostranslate.tags import translate_tags
from argostranslate.translate import ITranslation
from docx import Document
import sys
from argostranslatefiles.formats.abstract_xml import AbstractXml


class Docx(AbstractXml):
    supported_file_extensions = ['.docx']

    def translate_paragraphs(self,paragraphs,underlying_translation):
        for paragraph in paragraphs:
            sys.stderr.write(f'{paragraph.text}\n')

            if len(paragraph.runs) == 1:
                paragraph.runs[0].text = underlying_translation.translate(paragraph.runs[0].text)
            elif len(paragraph.runs) == 0:
                continue
            else:
                first_run_style = paragraph.runs[0].style
                #get the dominant run in the paragraph
                dominant_run = max(paragraph.runs, key=lambda x: len(x.text))
                sys.stderr.write(f'dominant run: {dominant_run.text}.\n')
                dom_font = dominant_run.font.name
                dom_bold = dominant_run.font.bold
                dom_italic = dominant_run.font.italic
                dom_underline = dominant_run.font.underline

                paragraph.text = underlying_translation.translate(paragraph.text)
                paragraph.runs[0].font.name = dom_font
                paragraph.runs[0].font.bold = dom_bold
                paragraph.runs[0].font.italic = dom_italic
                paragraph.runs[0].font.underline = dom_underline



    def translate(self, underlying_translation: ITranslation, file_path: str):
        outzip_path = self.get_output_path(underlying_translation, file_path)
        document = Document(file_path)

        self.translate_paragraphs(document.paragraphs,underlying_translation)
        
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.translate_paragraphs(cell.paragraphs,underlying_translation)

        document.save(outzip_path)

        return outzip_path
